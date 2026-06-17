import streamlit as st
import pandas as pd
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import plotly.graph_objects as go

def generate_word_report(S, Z, M, H, W_vid, P, ITS):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc = Document()

    def grade_text(v):
        if v >= 80:   return "Хорошо (≥ 80 %)"
        elif v >= 50: return "Удовлетворительно (50–80 %)"
        else:         return "Неудовлетворительно (< 50 %)"

    def _run(run, size=14, bold=False, italic=False):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    def _para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=False,
              sb=0, sa=4, ls=1.5):
        p.alignment = align
        pf = p.paragraph_format
        pf.first_line_indent = Cm(1.25) if indent else Cm(0)
        pf.space_before = Pt(sb)
        pf.space_after  = Pt(sa)
        pf.line_spacing = ls  # float = multiple; Pt(...) = exact

    def add_p(text="", size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              indent=False, bold=False, italic=False, sb=0, sa=4, ls=1.5):
        p = doc.add_paragraph()
        _para(p, align=align, indent=indent, sb=sb, sa=sa, ls=ls)
        if text:
            r = p.add_run(text)
            _run(r, size=size, bold=bold, italic=italic)
        return p

    def add_h(text, level=1):
        p = doc.add_paragraph()
        _para(p, align=WD_ALIGN_PARAGRAPH.LEFT, sb=8, sa=4)
        r = p.add_run(text)
        _run(r, size=14, bold=True)
        return p

    # ── Таблица: ячейки ──
    COL_W = [Cm(7.5), Cm(1.8), Cm(2.5), Cm(5.2)]  # итого ≈ 17 cm

    def set_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, size=12):
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = align
        pf = para.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(3)
        pf.space_after  = Pt(3)
        pf.line_spacing = Pt(13)  # точный одинарный
        r = para.add_run(text)
        _run(r, size=size, bold=bold)

    def thick_cell(cell, idx):
        """Жирные рамки 1,5 пт + ширина столбца."""
        cell.width = COL_W[idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        borders = OxmlElement('w:tcBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '12')   # 12/8 = 1,5 пт
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)
        tcPr.append(borders)

    # ═══════════════════════════════════════════════
    # Заголовок
    # ═══════════════════════════════════════════════
    p = doc.add_paragraph()
    _para(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=2)
    _run(p.add_run("Оценка эффективности интеллектуальной транспортной системы"),
         size=14, bold=True)

    add_p("По методологии Евстигнеева И.А. | БГТУ им. В.Г. Шухова, 2026",
          size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, sa=2)

    add_p(f"Дата формирования отчёта: {datetime.now().strftime('%d.%m.%Y  %H:%M')}",
          size=12, align=WD_ALIGN_PARAGRAPH.LEFT, sa=6)

    # ═══════════════════════════════════════════════
    # Раздел 1: Таблица
    # ═══════════════════════════════════════════════
    add_h("1 Результаты по подсистемам")

    add_p("Таблица 1 — Результаты по подсистемам ИТС",
          align=WD_ALIGN_PARAGRAPH.LEFT, sb=2, sa=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Фиксированная ширина таблицы (17 cm)
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(17 * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    hdr = table.rows[0].cells
    for i, txt in enumerate(["Подсистема", "Показатель", "Значение, %", "Оценка"]):
        set_cell(hdr[i], txt, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        thick_cell(hdr[i], i)

    rows_data = [
        ("I.   Светофорное управление",          "S", S),
        ("II.  Безопасность дорожного движения", "Z", Z),
        ("III. Мониторинг транспортного потока", "M", M),
        ("IV.  Метеомониторинг",                 "H", H),
        ("V.   Видеонаблюдение и инциденты",     "W", W_vid),
        ("VI.  Общественный транспорт",           "P", P),
    ]
    for name, sym, val in rows_data:
        row = table.add_row().cells
        set_cell(row[0], name)
        set_cell(row[1], sym,            align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row[2], f"{val:.2f}",   align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(row[3], grade_text(val))
        for i, cell in enumerate(row):
            thick_cell(cell, i)

    # ═══════════════════════════════════════════════
    # Раздел 2: Формула
    # ═══════════════════════════════════════════════
    add_h("2 Сводная оценка эффективности ИТС")

    p_f = add_p(sb=2, sa=1)
    _run(p_f.add_run(
        "ИТСэф = 0,2·S + 0,2·Z + 0,1·M + 0,1·H + 0,2·W + 0,2·P"
    ), size=13)

    # Пояснения — точный одинарный интервал (ГОСТ 6.10: «где» без двоеточия)
    gde = [
        "где  S — показатель светофорного управления, %;",
        "        Z — показатель безопасности дорожного движения, %;",
        "        M — показатель мониторинга транспортного потока, %;",
        "        H — показатель метеомониторинга, %;",
        "        W — показатель видеонаблюдения и выявления инцидентов, %;",
        "        P — показатель организации движения общественного транспорта, %.",
    ]
    for i, line in enumerate(gde):
        p = doc.add_paragraph()
        _para(p, align=WD_ALIGN_PARAGRAPH.LEFT,
              sb=3 if i == 0 else 0,
              sa=3 if i == len(gde)-1 else 0,
              ls=Pt(13))          # точный одинарный
        _run(p.add_run(line), size=12)

    add_p(
        f"ИТСэф = 0,2×{S:.2f} + 0,2×{Z:.2f} + 0,1×{M:.2f} + "
        f"0,1×{H:.2f} + 0,2×{W_vid:.2f} + 0,2×{P:.2f} = {ITS:.2f} %",
        size=13, sb=2, sa=4
    )

    # ═══════════════════════════════════════════════
    # Раздел 3: Итоговая оценка
    # ═══════════════════════════════════════════════
    add_h("3 Итоговая оценка")

    p_res = add_p(sb=0, sa=0)
    _run(p_res.add_run(f"ИТСэф = {ITS:.2f} % — {grade_text(ITS)}"),
         size=14, bold=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_pdf_report(S, Z, M, H, W_vid, P, ITS):
    import os
    from fpdf import FPDF

    def grade_text(v):
        if v >= 80:   return "Хорошо (>= 80 %)"
        elif v >= 50: return "Удовлетворительно (50-80 %)"
        else:         return "Неудовлетворительно (< 50 %)"

    def find_font(bold=False):
        variants = ["DejaVuSans-Bold.ttf", "DejaVuSans.ttf"] if bold else ["DejaVuSans.ttf"]
        linux_base = "/usr/share/fonts/truetype/dejavu/"
        win_base   = "C:\\Windows\\Fonts\\"
        win_names  = (["arialbd.ttf"] if bold else ["arial.ttf"])
        for name in variants:
            p = linux_base + name
            if os.path.exists(p): return p
        for name in win_names:
            p = win_base + name
            if os.path.exists(p): return p
        return None

    font_reg  = find_font(bold=False)
    font_bold = find_font(bold=True)

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(left=25, top=20, right=15)
    pdf.add_page()

    if font_reg:
        pdf.add_font("Main",  "",  font_reg)
        pdf.add_font("Main",  "B", font_bold or font_reg)
        fname = "Main"
    else:
        fname = "Helvetica"   # ASCII-only fallback

    # ── Заголовок ──
    pdf.set_font(fname, "B", 14)
    pdf.multi_cell(0, 8,
        "Оценka effektivnosti intellektual'noj transportnoj sistemy"
        if fname == "Helvetica" else
        "Оценка эффективности интеллектуальной транспортной системы",
        align="C")
    pdf.ln(2)
    pdf.set_font(fname, "", 11)
    pdf.cell(0, 6,
        "Po metodologii Evstigneeva I.A. | BGTU im. V.G. Shukhova, 2026"
        if fname == "Helvetica" else
        "По методологии Евстигнеева И.А. | БГТУ им. В.Г. Шухова, 2026",
        align="C")
    pdf.ln(8)
    pdf.set_font(fname, "", 12)
    pdf.cell(0, 6, f"Дата: {datetime.now().strftime('%d.%m.%Y  %H:%M')}")
    pdf.ln(10)

    # ── Раздел 1: Таблица ──
    pdf.set_font(fname, "B", 13)
    pdf.cell(0, 7, "1  Результаты по подсистемам")
    pdf.ln(5)
    pdf.set_font(fname, "", 11)
    pdf.cell(0, 5, "Таблица 1 — Результаты по подсистемам ИТС")
    pdf.ln(3)

    # Ширины: 170 мм = 25 левое поле; A4 = 210; правое = 15 → контент 170 мм
    cw = [72, 18, 26, 54]
    pdf.set_font(fname, "B", 11)
    for txt, w in zip(["Подсистема", "Показ.", "Значение, %", "Оценка"], cw):
        pdf.cell(w, 8, txt, border=1, align="C")
    pdf.ln()

    pdf.set_font(fname, "", 11)
    rows = [
        ("I.   Светофорное управление",          "S",  S),
        ("II.  Безопасность дорожного движения", "Z",  Z),
        ("III. Мониторинг транспортного потока", "M",  M),
        ("IV.  Метеомониторинг",                 "H",  H),
        ("V.   Видеонаблюдение и инциденты",     "W",  W_vid),
        ("VI.  Общественный транспорт",           "P",  P),
    ]
    for name, sym, val in rows:
        pdf.cell(cw[0], 7, name, border=1)
        pdf.cell(cw[1], 7, sym,          border=1, align="C")
        pdf.cell(cw[2], 7, f"{val:.2f}", border=1, align="C")
        pdf.cell(cw[3], 7, grade_text(val), border=1)
        pdf.ln()

    pdf.ln(8)

    # ── Раздел 2: Формула ──
    pdf.set_font(fname, "B", 13)
    pdf.cell(0, 7, "2  Сводная оценка эффективности ИТС")
    pdf.ln(5)
    pdf.set_font(fname, "", 12)
    pdf.cell(0, 6, "ИТСэф = 0,2·S + 0,2·Z + 0,1·M + 0,1·H + 0,2·W + 0,2·P")
    pdf.ln(5)

    pdf.set_font(fname, "", 11)
    for line in [
        "где  S — показатель светофорного управления, %;",
        "       Z — показатель безопасности дорожного движения, %;",
        "       M — показатель мониторинга транспортного потока, %;",
        "       H — показатель метеомониторинга, %;",
        "       W — показатель видеонаблюдения и выявления инцидентов, %;",
        "       P — показатель организации движения общественного транспорта, %.",
    ]:
        pdf.cell(0, 5, line)
        pdf.ln()

    pdf.ln(4)
    pdf.set_font(fname, "", 12)
    pdf.multi_cell(0, 6,
        f"ИТСэф = 0,2×{S:.2f} + 0,2×{Z:.2f} + 0,1×{M:.2f} + "
        f"0,1×{H:.2f} + 0,2×{W_vid:.2f} + 0,2×{P:.2f} = {ITS:.2f} %")
    pdf.ln(8)

    # ── Раздел 3: Итог ──
    pdf.set_font(fname, "B", 13)
    pdf.cell(0, 7, "3  Итоговая оценка")
    pdf.ln(5)
    pdf.set_font(fname, "B", 13)
    pdf.cell(0, 8, f"ИТСэф = {ITS:.2f} %  —  {grade_text(ITS)}")

    buf = io.BytesIO(pdf.output())
    buf.seek(0)
    return buf


st.set_page_config(
    page_title="Эффективность ИТС",
    page_icon="🚦",
    layout="wide"
)

st.markdown("""
<style>
/* Вкладки «📋 Итог» — всегда последние в группе → жирный шрифт */
[data-baseweb="tab-list"] button[role="tab"]:last-child p {
    font-weight: 700 !important;
}
/* Компактный сайдбар */
section[data-testid="stSidebar"] .stMarkdown p,
section[data-testid="stSidebar"] .stCaption p {
    margin-bottom: 2px !important;
    line-height: 1.3 !important;
}
section[data-testid="stSidebar"] h3 {
    margin-top: 4px !important;
    margin-bottom: 4px !important;
}
section[data-testid="stSidebar"] hr {
    margin-top: 6px !important;
    margin-bottom: 6px !important;
}
section[data-testid="stSidebar"] [data-testid="stProgressBar"] {
    margin-bottom: 2px !important;
}
</style>
""", unsafe_allow_html=True)

# ── Session state ──────────────────────────────────────────────
for k in ['S', 'Z', 'M', 'H', 'W_vid', 'P']:
    if k not in st.session_state:
        st.session_state[k] = None

# ── Helpers ────────────────────────────────────────────────────
def grade(v):
    if v >= 80:   return "🟢 Хорошо (≥ 80%)"
    elif v >= 50: return "🟡 Удовлетворительно (50–80%)"
    else:         return "🔴 Неудовлетворительно (< 50%)"

def save_button(label, key, value):
    if st.button(f"💾 Сохранить {label} для раздела VII", key=f"save_{key}"):
        st.session_state[key] = value
        st.toast(f"{label} = {value:.2f} % сохранён!", icon="✅")
        st.rerun()

def dt_block(key_prefix):
    with st.expander("Временной диапазон ΔT"):
        st.latex(r"\Delta T = T_1 - T_2")
        col1, col2 = st.columns(2)
        T2 = col1.number_input("T₂ — начало (порядковый номер суток)", 1, 365, 1, key=f"{key_prefix}_T2")
        T1 = col2.number_input("T₁ — конец (порядковый номер суток)",  1, 365, 90, key=f"{key_prefix}_T1")
        st.metric("ΔT", f"{T1 - T2} суток")

# ── Sidebar ────────────────────────────────────────────────────
with st.sidebar:
    st.title("🚦 Калькулятор ИТС")
    st.caption("Евстигнеев И.А. | БГТУ им. В.Г. Шухова, 2026")
    st.markdown("<div style='margin-bottom:8px'></div>", unsafe_allow_html=True)

    section = st.radio("Раздел:", [
        "I. Светофорное управление",
        "II. Безопасность дорожного движения",
        "III. Мониторинг транспортного потока",
        "IV. Метеомониторинг",
        "V. Видеонаблюдение и инциденты",
        "VI. Общественный транспорт",
        "VII. Сводная оценка ИТС",
    ])

    st.divider()
    saved_count = sum(1 for k in ['S', 'Z', 'M', 'H', 'W_vid', 'P']
                      if st.session_state[k] is not None)
    st.subheader("Прогресс")
    st.progress(saved_count / 6)
    st.caption(f"Заполнено {saved_count} / 6 разделов")
    st.markdown("<div style='margin-bottom:8px'></div>", unsafe_allow_html=True)

    sidebar_labels = [
        ('S',     'S — Светофоры'),
        ('Z',     'Z — БДД'),
        ('M',     'M — Мониторинг'),
        ('H',     'H — Метео'),
        ('W_vid', 'W — Видео'),
        ('P',     'P — НГПТ'),
    ]
    for k, lbl in sidebar_labels:
        v = st.session_state[k]
        if v is not None:
            icon = "🟢" if v >= 80 else ("🟡" if v >= 50 else "🔴")
            st.write(f"{icon} **{lbl}:** {v:.2f} %")
        else:
            st.caption(f"⬜ {lbl}: не заполнен")

    st.divider()
    if st.button("Заполнить тестовыми данными", use_container_width=True, type="secondary"):
        test = {'S': 85.0, 'Z': 44.0, 'M': 78.0, 'H': 62.0, 'W_vid': 88.0, 'P': 71.0}
        for k, v in test.items():
            st.session_state[k] = v
        st.rerun()

    if not st.session_state.get("_confirm_reset"):
        if st.button("Сбросить все сохранённые результаты", use_container_width=True, type="secondary"):
            st.session_state["_confirm_reset"] = True
            st.rerun()
    else:
        st.warning("Все введённые данные будут удалены. Продолжить?")
        c1, c2 = st.columns(2)
        if c1.button("✅ Да", use_container_width=True, type="primary"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()
        if c2.button("❌ Отмена", use_container_width=True):
            st.session_state["_confirm_reset"] = False
            st.rerun()


# ══════════════════════════════════════════════════════════════════
# I. СВЕТОФОРНОЕ УПРАВЛЕНИЕ
# ══════════════════════════════════════════════════════════════════
if section.startswith("I. "):
    st.header("I. Эффективность светофорного управления")
    st.latex(r"S = 0{,}4\,S_1 + 0{,}2\,S_2 + 0{,}2\,S_3 + 0{,}2\,S_4")

    t1, t2, t3, t4, t5 = st.tabs([
        "S₁ — Уровень развития",
        "S₂ — Задержка ТС",
        "S₃ — Остановки ТС",
        "S₄ — Скорость ТС",
        "📋 Итог S",
    ])

    # ── S1 ──
    with t1:
        st.subheader("S₁ — Уровень развития подсистемы")
        st.latex(r"S_1 = (S_{1.1} + S_{1.2})\times 100\%")
        st.info("Максимум S₁ = 200 % (полное адаптивное И диспетчерское покрытие одновременно)")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**S₁.₁ — адаптивное управление**")
            st.latex(r"S_{1.1}=\dfrac{a}{c}")
            a  = st.number_input("a — светофоров с адаптивным управлением, шт.", 0, 100000, 30)
            c1 = st.number_input("c — всего светофорных объектов, шт.", 1, 100000, 100, key="c1")
            s11 = a / c1
            st.metric("S₁.₁", f"{s11:.4f}", help=f"{s11*100:.1f} %")
        with col2:
            st.markdown("**S₁.₂ — диспетчерское управление**")
            st.latex(r"S_{1.2}=\dfrac{b}{c}")
            b  = st.number_input("b — светофоров с диспетчерским управлением, шт.", 0, 100000, 60)
            c2 = st.number_input("c — всего светофорных объектов, шт.", 1, 100000, 100, key="c2")
            s12 = b / c2
            st.metric("S₁.₂", f"{s12:.4f}", help=f"{s12*100:.1f} %")
        s1 = (s11 + s12) * 100
        st.divider()
        st.metric("🎯 S₁", f"{s1:.2f} %",
                  help="Макс. = 200 % (полное адаптивное И диспетчерское покрытие). ≥ 160 % — хорошо; ≥ 100 % — удовлетворительно; < 100 % — неудовлетворительно. Формула (3).")
        g = "🟢 Хорошо" if s1 >= 160 else ("🟡 Удовлетворительно" if s1 >= 100 else "🔴 Неудовлетворительно")
        st.caption(g)

    # ── S2 ──
    with t2:
        st.subheader("S₂ — Изменение средней задержки ТС")
        st.latex(r"S_2 = \frac{\tau_{s2}}{\tau_{s1}}\times 100\%")
        st.info("S₂ < 100 % → задержка снизилась (хорошо)  |  S₂ > 100 % → задержка выросла (плохо)")

        mode2 = st.radio("Режим ввода:", ["Ввести τs напрямую", "Рассчитать τs по участкам дороги"], key="mode2")

        if mode2 == "Ввести τs напрямую":
            col1, col2 = st.columns(2)
            ts1 = col1.number_input("τs₁ — задержка предыдущий период, ч/км", 0.0001, 100.0, 0.025, format="%.4f")
            ts2 = col2.number_input("τs₂ — задержка текущий период, ч/км",    0.0,    100.0, 0.020, format="%.4f")
        else:
            st.markdown("#### Формула (7)")
            st.latex(r"\tau_s = \frac{\displaystyle\sum_{i=1}^{n} m_i\cdot\tau_i}{\displaystyle\sum_{i=1}^{n} m_i\cdot l_i}\,,\quad\text{ч/км}")
            st.caption("τᵢ = T − Tсв  (задержка = фактическое время − время свободного движения)")

            def calc_ts(period_lbl, key_pfx, def_T, def_Tsv):
                st.markdown(f"**{period_lbl}:**")
                n = st.number_input("Число участков", 1, 20, 2, key=f"{key_pfx}_n")
                num, den = 0.0, 0.0
                for i in range(n):
                    c1, c2, c3, c4 = st.columns(4)
                    mi  = c1.number_input(f"mᵢ (полос) [{i+1}]",     1,    20,   2,       key=f"{key_pfx}_m{i}")
                    T_i = c2.number_input(f"T (ч) [{i+1}]",          0.001, 10.0, def_T,   format="%.4f", key=f"{key_pfx}_T{i}")
                    Sv  = c3.number_input(f"Tсв (ч) [{i+1}]",        0.001, 10.0, def_Tsv, format="%.4f", key=f"{key_pfx}_Sv{i}")
                    li  = c4.number_input(f"l (км) [{i+1}]",         0.1,  500.0, 1.0,     format="%.2f",  key=f"{key_pfx}_l{i}")
                    tau_i = max(0.0, T_i - Sv)
                    num += mi * tau_i
                    den += mi * li
                ts = num / den if den > 0 else 0.0
                st.metric(f"τs ({period_lbl})", f"{ts:.4f} ч/км")
                return ts

            col1, col2 = st.columns(2)
            with col1: ts1 = calc_ts("Предыдущий период", "prev2", 0.050, 0.030)
            with col2: ts2 = calc_ts("Текущий период",    "curr2", 0.040, 0.030)

        s2 = (ts2 / ts1 * 100) if ts1 > 0 else 0.0
        st.divider()
        st.metric("🎯 S₂", f"{s2:.2f} %",
                  help="< 100 % — задержка снизилась (хорошо); > 100 % — задержка выросла (плохо). Формула (6).")
        if s2 <= 100:
            st.caption(f"🟢 Задержка снизилась на {100-s2:.1f} %")
        else:
            st.caption(f"🔴 Задержка выросла на {s2-100:.1f} %")

    # ── S3 ──
    with t3:
        st.subheader("S₃ — Изменение числа остановок ТС")
        mode3 = st.radio("Режим:", ["Ввести Nост напрямую", "Рассчитать по формуле Вебстера (14)"], key="mode3")

        if mode3 == "Ввести Nост напрямую":
            col1, col2 = st.columns(2)
            n1 = col1.number_input("Nост₁.ср — остановок предыдущий период", 0.001, 1e6, 2.5, format="%.3f")
            n2 = col2.number_input("Nост₂.ср — остановок текущий период",    0.0,   1e6, 2.0, format="%.3f")
        else:
            st.latex(r"N_{\text{ост}} \approx \sum_{i=1}^{n}\frac{N_i\cdot t_{\text{кр},i}^{2}}{2\cdot T_\text{ц}\cdot(T_\text{ц}-t_{\text{кр},i})}")
            Tc = st.number_input("Tц — длительность цикла светофорного регулирования, с", 30.0, 300.0, 90.0, key="Tc")

            def webster(pfx, lbl, Tc):
                st.markdown(f"**{lbl}:**")
                n_app = st.number_input("Число подходов к перекрёстку", 1, 12, 4, key=f"{pfx}_napp")
                total = 0.0
                for i in range(n_app):
                    c1, c2 = st.columns(2)
                    Ni  = c1.number_input(f"N{i+1} — интенсивность (ТС/ч)", 0.0, 10000.0, 500.0, key=f"{pfx}_N{i}")
                    tkr = c2.number_input(f"tкр,{i+1} — красный сигнал, с",  1.0, Tc-1.0,  45.0,  key=f"{pfx}_tkr{i}")
                    if Tc > tkr > 0:
                        total += (Ni * tkr**2) / (2 * Tc * (Tc - tkr))
                st.metric(f"Nост.ср ({lbl})", f"{total:.3f}")
                return total

            col1, col2 = st.columns(2)
            with col1: n1 = webster("wb_prev", "Предыдущий период", Tc)
            with col2: n2 = webster("wb_curr", "Текущий период",    Tc)

        st.divider()
        if n1 > 0:
            if n1 >= n2:
                st.latex(r"S_3 = \left(1 - \frac{N_{\text{ост}2}}{N_{\text{ост}1}}\right)\times 100\%")
                s3 = (1 - n2 / n1) * 100
                st.caption("Формула (11): остановок стало меньше или равно")
            else:
                st.latex(r"S_3 = \frac{N_{\text{ост}2}}{N_{\text{ост}1}}\times 100\%")
                s3 = (n2 / n1) * 100
                st.caption("Формула (12): остановок стало больше")
        else:
            s3 = 0.0
        st.metric("🎯 S₃", f"{s3:.2f} %",
                  help="≥ 80 % — хорошо; 50–80 % — удовлетворительно; < 50 % — неудовлетворительно. Формула (11) или (12).")
        st.caption(grade(s3))

    # ── S4 ──
    with t4:
        st.subheader("S₄ — Изменение средней скорости ТС")
        st.latex(r"S_4 = \frac{V_{t2}}{V_м}\times 100\%")
        col1, col2 = st.columns(2)
        vt2 = col1.number_input("Vt₂ — средняя скорость ТС текущий период, км/ч", 0.1, 500.0, 45.0)
        vm  = col2.number_input("Vм — максимальная разрешённая скорость, км/ч",   0.1, 500.0, 60.0)
        s4 = (vt2 / vm) * 100
        st.divider()
        st.metric("🎯 S₄", f"{s4:.2f} %",
                  help="≥ 80 % — хорошо; 50–80 % — удовлетворительно; < 50 % — неудовлетворительно. Формула (15).")
        st.caption(grade(s4))

    # ── Итог S ──
    with t5:
        st.subheader("Итоговый показатель S")
        st.latex(r"S = 0{,}4\,S_1 + 0{,}2\,S_2 + 0{,}2\,S_3 + 0{,}2\,S_4")
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("S₁", f"{s1:.2f} %")
        col2.metric("S₂", f"{s2:.2f} %")
        col3.metric("S₃", f"{s3:.2f} %")
        col4.metric("S₄", f"{s4:.2f} %")
        S = 0.4*s1 + 0.2*s2 + 0.2*s3 + 0.2*s4
        st.divider()
        st.markdown(f"S = 0,4 × {s1:.2f} + 0,2 × {s2:.2f} + 0,2 × {s3:.2f} + 0,2 × {s4:.2f} = **{S:.2f} %**")
        st.metric("🏆 S — Светофорное управление", f"{S:.2f} %",
                  help="Формула (2). Вес в ИТСэф: 0,20. ≥ 80 % — хорошо; 50–80 % — удовлетворительно; < 50 % — неудовлетворительно.")
        st.caption(grade(S))
        save_button("S", "S", S)

    dt_block("s")


# ══════════════════════════════════════════════════════════════════
# II. БЕЗОПАСНОСТЬ ДОРОЖНОГО ДВИЖЕНИЯ
# ══════════════════════════════════════════════════════════════════
elif section.startswith("II. "):
    st.header("II. Оценка безопасности дорожного движения")
    st.info(
        "В разделе сравниваются два подхода к оценке аварийности:\n\n"
        "**Z₁** — классический показатель по ОДМ 218.6.027-2017 (чем меньше — тем безопаснее)\n\n"
        "**Z₂** — авторская формула Евстигнеева с весами тяжести ДТП (чем больше — тем лучше, результат в %)\n\n"
        "Для раздела VII используется **Z₂**."
    )

    t1, t2, t3 = st.tabs([
        "Z₁ — Аварийность (ОДМ 218.6.027-2017)",
        "Z₂ — Аварийность с весами (формула Евстигнеева)",
        "📋 Сравнение показателей",
    ])

    with t1:
        st.subheader("Z₁ — Показатель относительной аварийности")
        st.latex(r"Z_{1} = \frac{n \cdot 10^{6}}{N \cdot L \cdot m \cdot 365}")
        col1, col2 = st.columns(2)
        n_dtp = col1.number_input("n — число ДТП за расчётный период", 0, 1000000, 15)
        N_int = col1.number_input("N — среднегодовая суточная интенсивность, авт./сут", 1, 10000000, 10000)
        L_len = col2.number_input("L — длина участка, км", 0.1, 100000.0, 10.0)
        m_yrs = col2.number_input("m — число лет в расчётном периоде (3–5)", 1, 10, 3)
        Z18 = (n_dtp * 1e6) / (N_int * L_len * m_yrs * 365)
        st.divider()
        st.metric("🎯 Z₁", f"{Z18:.4f}",
                  help="ОДМ 218.6.027-2017: дороги I кат. — не более 0,20; II кат. — не более 0,30; III кат. — не более 0,40. Чем ниже — тем безопаснее. Формула (18).")
        st.caption("ОДМ 218.6.027-2017: чем ниже значение — тем безопаснее участок. "
                   "Ориентир: дороги I кат. — не более 0,20; II кат. — не более 0,30; III кат. — не более 0,40")

    with t2:
        st.subheader("Z₂ — Показатель аварийности с весовыми коэффициентами")
        st.latex(r"Z_{2} = \left(1 - \frac{A + \alpha D + \beta W}{L \cdot N}\right)\times 100\%")
        st.info("Z → 100 % — максимальная безопасность  |  Z → 0 % или < 0 % — высокая аварийность")

        is_agglom = st.checkbox("Городская агломерация (раздельный учёт по типам дорог)")

        if is_agglom:
            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown("*Магистральные общегородские*")
                L1 = st.number_input("L₁, км", 0.0, 1e5, 50.0,  key="L1")
                D1 = st.number_input("D₁ — погибших", 0, 100000, 5,  key="D1")
                W1 = st.number_input("W₁ — раненых",  0, 100000, 20, key="W1")
                N1 = st.number_input("N₁, авт./сут",  0, 1000000, 15000, key="N1")
            with c2:
                st.markdown("*Магистральные районные*")
                L2 = st.number_input("L₂, км", 0.0, 1e5, 100.0, key="L2")
                D2 = st.number_input("D₂ — погибших", 0, 100000, 8,  key="D2")
                W2 = st.number_input("W₂ — раненых",  0, 100000, 35, key="W2")
                N2 = st.number_input("N₂, авт./сут",  0, 1000000, 8000, key="N2")
            with c3:
                st.markdown("*Местного значения*")
                L3 = st.number_input("L₃, км", 0.0, 1e5, 150.0, key="L3")
                D3 = st.number_input("D₃ — погибших", 0, 100000, 3,  key="D3")
                W3 = st.number_input("W₃ — раненых",  0, 100000, 15, key="W3")
                N3 = st.number_input("N₃, авт./сут",  0, 1000000, 3000, key="N3")
            A_dtp = st.number_input("A — всего ДТП за расчётный период", 0, 1000000, 50)
            n_types = st.number_input("n — число типов дорог в расчёте", 1, 3, 3, key="n_types")
            L = L1 + L2 + L3
            D = D1 + D2 + D3
            W_inj = W1 + W2 + W3
            N = (N1 + N2 + N3) / n_types
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("L, км", f"{L:.1f}"); c2.metric("D", D); c3.metric("W", W_inj); c4.metric("N, авт./сут", f"{N:.0f}")
        else:
            col1, col2 = st.columns(2)
            A_dtp = col1.number_input("A — число ДТП",           0, 1000000, 50)
            D     = col1.number_input("D — число погибших",      0, 100000,  10)
            W_inj = col1.number_input("W — число раненых",       0, 100000,  40)
            L     = col2.number_input("L — протяжённость, км",   0.1, 1e5,   100.0)
            N     = col2.number_input("N — суточная интенсивность, авт./сут", 1.0, 1e7, 10000.0)

        st.divider()
        alpha = st.number_input("α — вес погибших (рекомендовано: 5)", 0.0, 100.0, 5.0)
        beta  = st.number_input("β — вес раненых (рекомендовано: 1)",  0.0, 100.0, 1.0)
        denom = L * N
        numer = A_dtp + alpha * D + beta * W_inj
        Z19 = (1 - numer / denom) * 100 if denom > 0 else 0.0
        st.markdown(f"Z₂ = (1 − ({A_dtp} + {alpha}×{D} + {beta}×{W_inj}) / ({L}×{N:.0f})) × 100 % = **{Z19:.2f} %**")
        st.metric("🎯 Z₂", f"{Z19:.2f} %",
                  help="Чем ближе к 100 % — тем безопаснее. ≥ 80 % — хорошо; 50–80 % — удовлетворительно; < 50 % — неудовлетворительно. Формула (19) Евстигнеева.")
        st.caption(grade(Z19))

    with t3:
        st.subheader("Z — Безопасность дорожного движения")
        col1, col2 = st.columns(2)
        col1.metric("Z₁ — ОДМ (справочно)", f"{Z18:.4f}")
        col2.metric("Z₂ — С весами (основная)", f"{Z19:.2f} %")
        st.divider()
        st.markdown(f"Z = **{Z19:.2f} %** (по формуле с весовыми коэффициентами)")
        st.metric("🏆 Z — Безопасность дорожного движения", f"{Z19:.2f} %",
                  help="Формула (19). Вес в ИТСэф: 0,20. Используется Z₂ с весовыми коэффициентами тяжести ДТП (α=5 погибших, β=1 раненых).")
        st.caption(grade(Z19))
        save_button("Z", "Z", Z19)

    dt_block("z")


# ══════════════════════════════════════════════════════════════════
# III. МОНИТОРИНГ ТРАНСПОРТНОГО ПОТОКА
# ══════════════════════════════════════════════════════════════════
elif section.startswith("III. "):
    st.header("III. Эффективность подсистемы мониторинга транспортного потока")
    st.latex(r"M = 0{,}3\,M_1 + 0{,}3\,M_2 + 0{,}1\,M_3 + 0{,}1\,M_4 + 0{,}2\,M_5")

    t1, t2, t3, t4, t5, t6 = st.tabs(["M₁ — Качество данных", "M₂ — Достоверность",
                                        "M₃ — Охват", "M₄ — Аналитика", "M₅ — Оперативность", "📋 Итог M"])
    with t1:
        st.subheader("M₁ — Качество и регулярность данных")
        st.latex(r"M_1 = \frac{M_{1.1}+M_{1.2}}{n}")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**M₁.₁ — данные с пунктов учёта**")
            st.latex(r"M_{1.1}=\frac{a}{b}\times 100\%")
            a_m  = st.number_input("a — исправных датчиков, шт.", 0, 100000, 45)
            b_m  = st.number_input("b — всего датчиков, шт.",     1, 100000, 50, key="b_m11")
            m11 = (a_m / b_m) * 100
            st.metric("M₁.₁", f"{m11:.2f} %")
        with col2:
            st.markdown("**M₁.₂ — данные от бортовых систем (ГЛОНАСС/GPS)**")
            st.latex(r"M_{1.2}=\frac{p}{d}\times 100\%")
            p_m = st.number_input("p — суток с почасовой телеметрией", 0, 365, 85)
            d_m = st.number_input("d — контрольных суток",             1, 365, 90, key="d_m12")
            m12 = (p_m / d_m) * 100
            st.metric("M₁.₂", f"{m12:.2f} %")
        n_meth = st.number_input("n — число методов сбора данных", 1, 10, 2, key="n_meth")
        m1 = (m11 + m12) / n_meth
        st.divider(); st.metric("🎯 M₁", f"{m1:.2f} %"); st.caption(grade(m1))

    with t2:
        st.subheader("M₂ — Достоверность данных")
        st.latex(r"M_2=\frac{c}{d}\times 100\%")
        col1, col2 = st.columns(2)
        c_m  = col1.number_input("c — суток без аномальных данных", 0, 365, 88)
        d_m2 = col2.number_input("d — контрольных суток",           1, 365, 90, key="d_m2")
        m2 = (c_m / d_m2) * 100
        st.divider(); st.metric("🎯 M₂", f"{m2:.2f} %"); st.caption(grade(m2))

    with t3:
        st.subheader("M₃ — Полнота охвата мониторингом")
        st.latex(r"M_3=\frac{b}{f}\times 100\%")
        col1, col2 = st.columns(2)
        b_m3 = col1.number_input("b — точек мониторинга фактически, шт.", 0, 100000, 45)
        f_m3 = col2.number_input("f — точек мониторинга требуется, шт.",  1, 100000, 50)
        m3 = (b_m3 / f_m3) * 100
        st.divider(); st.metric("🎯 M₃", f"{m3:.2f} %"); st.caption(grade(m3))

    with t4:
        st.subheader("M₄ — Полнота аналитической обработки")
        st.latex(r"M_4=\frac{g}{d}\times 100\%")
        st.caption("g — суток, когда измерялись: интенсивность, состав потока И рассчитывались: скорость, плотность, пропускная способность, эффективность ОДД")
        col1, col2 = st.columns(2)
        g_m  = col1.number_input("g — суток с полной аналитикой", 0, 365, 80)
        d_m4 = col2.number_input("d — контрольных суток",         1, 365, 90, key="d_m4")
        m4 = (g_m / d_m4) * 100
        st.divider(); st.metric("🎯 M₄", f"{m4:.2f} %"); st.caption(grade(m4))

    with t5:
        st.subheader("M₅ — Оперативность передачи данных в АСУДД")
        st.latex(r"M_5=\frac{n}{d}\times 100\%")
        col1, col2 = st.columns(2)
        n_m5 = col1.number_input("n — суток с оперативной передачей всех изменений", 0, 365, 87)
        d_m5 = col2.number_input("d — контрольных суток",                            1, 365, 90, key="d_m5")
        m5 = (n_m5 / d_m5) * 100
        st.divider(); st.metric("🎯 M₅", f"{m5:.2f} %"); st.caption(grade(m5))

    with t6:
        st.subheader("Итоговый показатель M")
        st.latex(r"M = 0{,}3\,M_1 + 0{,}3\,M_2 + 0{,}1\,M_3 + 0{,}1\,M_4 + 0{,}2\,M_5")
        cols = st.columns(5)
        for i, (v, n) in enumerate([(m1,"M₁"),(m2,"M₂"),(m3,"M₃"),(m4,"M₄"),(m5,"M₅")]):
            cols[i].metric(n, f"{v:.2f} %")
        M = 0.3*m1 + 0.3*m2 + 0.1*m3 + 0.1*m4 + 0.2*m5
        st.divider()
        st.markdown(f"M = 0,3×{m1:.2f} + 0,3×{m2:.2f} + 0,1×{m3:.2f} + 0,1×{m4:.2f} + 0,2×{m5:.2f} = **{M:.2f} %**")
        st.metric("🏆 M — Мониторинг транспортного потока", f"{M:.2f} %",
                  help="Формула (25). Вес в ИТСэф: 0,10. ≥ 80 % — хорошо; 50–80 % — удовлетворительно; < 50 % — неудовлетворительно.")
        st.caption(grade(M))
        save_button("M", "M", M)

    dt_block("m")


# ══════════════════════════════════════════════════════════════════
# IV. МЕТЕОМОНИТОРИНГ
# ══════════════════════════════════════════════════════════════════
elif section.startswith("IV. "):
    st.header("IV. Эффективность подсистемы метеомониторинга")
    st.latex(r"H = 0{,}3\,H_1 + 0{,}1\,H_2 + 0{,}1\,H_3 + 0{,}1\,H_4 + 0{,}2\,H_5 + 0{,}2\,H_6")

    t1,t2,t3,t4,t5,t6,t7 = st.tabs(["H₁ — Качество данных","H₂ — Охват","H₃ — Аналитика","H₄ — Интеграция","H₅ — Эксплуатац. службы","H₆ — Информирование","📋 Итог H"])

    with t1:
        st.subheader("H₁ — Качество метеорологической информации")
        st.latex(r"H_1=\frac{a}{b}\times 100\%")
        col1,col2=st.columns(2)
        a_h=col1.number_input("a — работоспособных АДМС, шт.",1,100000,9)
        b_h=col2.number_input("b — всего АДМС, шт.",           1,100000,10)
        h1=(a_h/b_h)*100; st.metric("🎯 H₁",f"{h1:.2f} %"); st.caption(grade(h1))

    with t2:
        st.subheader("H₂ — Полнота охвата метеонаблюдением")
        st.latex(r"H_2=\frac{c}{d}\times 100\%")
        col1,col2=st.columns(2)
        c_h=col1.number_input("c — мест дислокации АДМС фактически", 0,100000,8)
        d_h=col2.number_input("d — мест, где необходим метеомониторинг",1,100000,10)
        h2=(c_h/d_h)*100; st.metric("🎯 H₂",f"{h2:.2f} %"); st.caption(grade(h2))

    with t3:
        st.subheader("H₃ — Качество аналитической обработки метеоданных")
        st.latex(r"H_3=\frac{f}{k}\times 100\%")
        st.caption("Условие: расхождение прогноза с реальными погодными условиями ≤ 30 % (ГОСТ Р 71094-2024)")
        col1,col2=st.columns(2)
        f_h =col1.number_input("f — суток с метеопрогнозом по базовым параметрам",0,365,88)
        k_h3=col2.number_input("k — контрольных суток",1,365,90,key="k_h3")
        h3=(f_h/k_h3)*100; st.metric("🎯 H₃",f"{h3:.2f} %"); st.caption(grade(h3))

    with t4:
        st.subheader("H₄ — Интеграция с внешними информационными системами")
        st.latex(r"H_4=\frac{j}{k}\times 100\%")
        col1,col2=st.columns(2)
        j_h =col1.number_input("j — суток с передачей данных во внешние системы",0,365,85)
        k_h4=col2.number_input("k — контрольных суток",1,365,90,key="k_h4")
        h4=(j_h/k_h4)*100; st.metric("🎯 H₄",f"{h4:.2f} %"); st.caption(grade(h4))

    with t5:
        st.subheader("H₅ — Своевременность информирования эксплуатационных служб")
        st.latex(r"H_5=\frac{l}{k}\times 100\%")
        st.caption("Включает: прогноз на 48 ч, штормовые предупреждения, производственно-технологические предупреждения ≥ 4 ч (ГОСТ Р 71094-2024)")
        col1,col2=st.columns(2)
        l_h =col1.number_input("l — суток со своевременным информированием служб",0,365,87)
        k_h5=col2.number_input("k — контрольных суток",1,365,90,key="k_h5")
        h5=(l_h/k_h5)*100; st.metric("🎯 H₅",f"{h5:.2f} %"); st.caption(grade(h5))

    with t6:
        st.subheader("H₆ — Своевременность информирования участников дорожного движения")
        st.latex(r"H_6=\frac{n}{k}\times 100\%")
        col1,col2=st.columns(2)
        n_h =col1.number_input("n — суток со своевременным информированием ДД",0,365,86)
        k_h6=col2.number_input("k — контрольных суток",1,365,90,key="k_h6")
        h6=(n_h/k_h6)*100; st.metric("🎯 H₆",f"{h6:.2f} %"); st.caption(grade(h6))

    with t7:
        st.subheader("Итоговый показатель H")
        st.latex(r"H = 0{,}3\,H_1+0{,}1\,H_2+0{,}1\,H_3+0{,}1\,H_4+0{,}2\,H_5+0{,}2\,H_6")
        cols=st.columns(6)
        for i,(v,n) in enumerate([(h1,"H₁"),(h2,"H₂"),(h3,"H₃"),(h4,"H₄"),(h5,"H₅"),(h6,"H₆")]):
            cols[i].metric(n,f"{v:.2f} %")
        H=0.3*h1+0.1*h2+0.1*h3+0.1*h4+0.2*h5+0.2*h6
        st.divider()
        st.markdown(f"H = 0,3×{h1:.2f} + 0,1×{h2:.2f} + 0,1×{h3:.2f} + 0,1×{h4:.2f} + 0,2×{h5:.2f} + 0,2×{h6:.2f} = **{H:.2f} %**")
        st.metric("🏆 H — Метеомониторинг", f"{H:.2f} %",
                  help="Формула (34). Вес в ИТСэф: 0,10. ≥ 80 % — хорошо; 50–80 % — удовлетворительно; < 50 % — неудовлетворительно.")
        st.caption(grade(H))
        save_button("H","H",H)

    dt_block("h")


# ══════════════════════════════════════════════════════════════════
# V. ВИДЕОНАБЛЮДЕНИЕ И ИНЦИДЕНТЫ
# ══════════════════════════════════════════════════════════════════
elif section.startswith("V. "):
    st.header("V. Эффективность подсистемы видеонаблюдения и выявления инцидентов")
    st.latex(r"W = 0{,}3\,W_1 + 0{,}2\,W_2 + 0{,}2\,W_3 + 0{,}3\,W_4")

    t1,t2,t3,t4,t5 = st.tabs(["W₁ — Качество видео","W₂ — Охват","W₃ — Аналитика","W₄ — Реагирование","📋 Итог W"])

    with t1:
        st.subheader("W₁ — Качество видеоинформации")
        st.latex(r"W_1=\frac{a}{b}\times 100\%")
        col1,col2=st.columns(2)
        a_w=col1.number_input("a — работоспособных видеокамер, шт.",0,1000000,95)
        b_w=col2.number_input("b — всего видеокамер, шт.",           1,1000000,100)
        w1=(a_w/b_w)*100
        st.caption("Условие: качество изображения соответствует проектной документации и паспорту оборудования")
        st.metric("🎯 W₁",f"{w1:.2f} %"); st.caption(grade(w1))

    with t2:
        st.subheader("W₂ — Полнота охвата видеонаблюдением")
        st.latex(r"W_2=\frac{c}{d}\times 100\%")
        col1,col2=st.columns(2)
        c_w=col1.number_input("c — мест дислокации видеокамер",       0,1000000,80)
        d_w=col2.number_input("d — мест, где необходимо видеонаблюдение",1,1000000,100)
        w2=(c_w/d_w)*100; st.metric("🎯 W₂",f"{w2:.2f} %"); st.caption(grade(w2))

    with t3:
        st.subheader("W₃ — Аналитическая обработка видео")
        st.latex(r"W_3=\frac{e}{b}\times 100\%")
        col1,col2=st.columns(2)
        e_w =col1.number_input("e — камер с видеоаналитикой, шт.",0,1000000,60)
        b_w3=col2.number_input("b — всего видеокамер, шт.",       1,1000000,100,key="b_w3")
        w3=(e_w/b_w3)*100
        st.caption("Аналитика должна включать: остановившиеся ТС, встречное движение, резкое торможение, пешеходы, упавшие объекты, классификация ТС")
        st.metric("🎯 W₃",f"{w3:.2f} %"); st.caption(grade(w3))

    with t4:
        st.subheader("W₄ — Готовность к ликвидации последствий инцидентов")
        st.latex(r"W_4=\frac{j}{k}\times 100\%")
        col1,col2=st.columns(2)
        j_w=col1.number_input("j — инцидентов с нормативной ликвидацией последствий",0,1000000,45)
        k_w=col2.number_input("k — всего выявленных инцидентов",                      1,1000000,50)
        w4=(j_w/k_w)*100
        st.caption("Условие: разработаны и применяются нормативные документы о сроках ликвидации инцидентов")
        st.metric("🎯 W₄",f"{w4:.2f} %"); st.caption(grade(w4))

    with t5:
        st.subheader("Итоговый показатель W — Видеонаблюдение")
        st.latex(r"W = 0{,}3\,W_1+0{,}2\,W_2+0{,}2\,W_3+0{,}3\,W_4")
        col1,col2,col3,col4=st.columns(4)
        col1.metric("W₁",f"{w1:.2f} %"); col2.metric("W₂",f"{w2:.2f} %")
        col3.metric("W₃",f"{w3:.2f} %"); col4.metric("W₄",f"{w4:.2f} %")
        W_vid=0.3*w1+0.2*w2+0.2*w3+0.3*w4
        st.divider()
        st.markdown(f"W = 0,3×{w1:.2f} + 0,2×{w2:.2f} + 0,2×{w3:.2f} + 0,3×{w4:.2f} = **{W_vid:.2f} %**")
        st.metric("🏆 W — Видеонаблюдение и инциденты", f"{W_vid:.2f} %",
                  help="Формула (42). Вес в ИТСэф: 0,20. ≥ 80 % — хорошо; 50–80 % — удовлетворительно; < 50 % — неудовлетворительно.")
        st.caption(grade(W_vid))
        save_button("W","W_vid",W_vid)

    dt_block("w")


# ══════════════════════════════════════════════════════════════════
# VI. ОБЩЕСТВЕННЫЙ ТРАНСПОРТ
# ══════════════════════════════════════════════════════════════════
elif section.startswith("VI. "):
    st.header("VI. Эффективность организации движения общественного транспорта")
    st.caption("⚠️  Применяется только для городских агломераций")
    st.latex(r"P = 0{,}4\,P_1 + 0{,}4\,P_2 + 0{,}2\,P_3")

    t1,t2,t3,t4 = st.tabs(["P₁ — Скорость НГПТ","P₂ — Пунктуальность","P₃ — Пассажиропоток","📋 Итог P"])

    with t1:
        st.subheader("P₁ — Изменение средней скорости НГПТ")
        st.latex(r"P_1=\frac{V_2}{V_м}\times 100\%")
        col1,col2=st.columns(2)
        V2=col1.number_input("V₂ — средняя скорость НГПТ текущий период, км/ч",0.1,300.0,20.0)
        Vm=col2.number_input("Vм — максимальная скорость НГПТ на участке, км/ч",0.1,300.0,40.0)
        p1=(V2/Vm)*100
        st.metric("🎯 P₁",f"{p1:.2f} %"); st.caption(grade(p1))
        with st.expander("Дополнительный показатель P₁к (для местного уровня)"):
            st.latex(r"P_{1к}=\frac{V_2}{V_1}\times 100\%")
            V1=st.number_input("V₁ — скорость за предыдущий период, км/ч",0.1,300.0,18.0)
            st.metric("P₁к",f"{(V2/V1)*100:.2f} %")

    with t2:
        st.subheader("P₂ — Пунктуальность (соблюдение графика перевозок)")
        st.latex(r"P_2=\frac{r_2}{r_м}\times 100\%")
        st.markdown("**Расчёт r₂ (формула 52):**")
        st.latex(r"r_2=\frac{a}{b}")
        col1,col2=st.columns(2)
        a_p=col1.number_input("a — рейсов по расписанию", 0,10000000,850)
        b_p=col2.number_input("b — всего рейсов НГПТ",    1,10000000,1000)
        r2=a_p/b_p
        st.metric("r₂",f"{r2:.4f}")
        rm=st.number_input("rм — максимально возможное соотношение (обычно = 1,0)",0.01,1.0,1.0,format="%.2f")
        p2=(r2/rm)*100
        st.metric("🎯 P₂",f"{p2:.2f} %"); st.caption(grade(p2))
        with st.expander("Дополнительный показатель P₂к (для местного уровня)"):
            st.latex(r"P_{2к}=\frac{r_2}{r_1}\times 100\%")
            r1=st.number_input("r₁ — соотношение за предыдущий период",0.01,1.0,0.80,format="%.4f")
            st.metric("P₂к",f"{(r2/r1)*100:.2f} %")

    with t3:
        st.subheader("P₃ — Изменение величины пассажиропотока НГПТ")
        col1,col2=st.columns(2)
        e2=col1.number_input("e₂ — пассажиропоток текущий период, чел.",0,100000000,95000)
        eb=col2.number_input("eб — базовый пассажиропоток, чел.",        1,100000000,100000)
        if eb >= e2:
            st.latex(r"P_3=\frac{e_2}{e_б}\times 100\%")
            p3=(e2/eb)*100
            st.caption("Формула (54): eб ≥ e₂ — пассажиропоток на уровне базового или ниже")
        else:
            st.latex(r"P_3=\left(1-\frac{e_2}{e_б}\right)\times 100\%")
            p3=(1-e2/eb)*100
            st.caption("Формула (55): eб < e₂ — пассажиропоток выше базового")
            st.warning("Отрицательный P₃ означает: провозная способность не справляется с пассажиропотоком (переполненность)")
        st.metric("🎯 P₃",f"{p3:.2f} %"); st.caption(grade(p3))
        with st.expander("Дополнительный показатель P₃к (для местного уровня)"):
            st.latex(r"P_{3к}=\frac{e_2}{e_1}\times 100\%")
            e1=st.number_input("e₁ — пассажиропоток предыдущий период, чел.",1,100000000,90000)
            st.metric("P₃к",f"{(e2/e1)*100:.2f} %")

    with t4:
        st.subheader("Итоговый показатель P — Общественный транспорт")
        st.latex(r"P = 0{,}4\,P_1+0{,}4\,P_2+0{,}2\,P_3")
        col1,col2,col3=st.columns(3)
        col1.metric("P₁",f"{p1:.2f} %"); col2.metric("P₂",f"{p2:.2f} %"); col3.metric("P₃",f"{p3:.2f} %")
        P=0.4*p1+0.4*p2+0.2*p3
        st.divider()
        st.markdown(f"P = 0,4×{p1:.2f} + 0,4×{p2:.2f} + 0,2×{p3:.2f} = **{P:.2f} %**")
        st.metric("🏆 P — Общественный транспорт", f"{P:.2f} %",
                  help="Формула (48). Вес в ИТСэф: 0,20. Применяется только для городских агломераций. ≥ 80 % — хорошо; 50–80 % — удовлетворительно.")
        st.caption(grade(P))
        save_button("P","P",P)

    dt_block("p")


# ══════════════════════════════════════════════════════════════════
# VII. СВОДНАЯ ОЦЕНКА ИТС
# ══════════════════════════════════════════════════════════════════
elif section.startswith("VII. "):
    st.header("VII. Сводная оценка эффективности ИТС")
    st.latex(r"\text{ИТС}_{\text{эф}} = 0{,}2\,S + 0{,}2\,Z + 0{,}1\,M + 0{,}1\,H + 0{,}2\,W + 0{,}2\,P")
    st.info("Введите значения вручную или используйте кнопки «Сохранить» в разделах I–VI — значения подставятся автоматически.")

    col1, col2 = st.columns(2)
    with col1:
        S     = st.number_input("S — Светофорное управление, %",          -500.0, 500.0, float(st.session_state['S']     or 75.0), format="%.2f", key="vS")
        Z     = st.number_input("Z — Безопасность дорожного движения, %", -500.0, 100.0, float(st.session_state['Z']     or 70.0), format="%.2f", key="vZ")
        M     = st.number_input("M — Мониторинг транспортного потока, %",    0.0, 100.0, float(st.session_state['M']     or 85.0), format="%.2f", key="vM")
    with col2:
        H     = st.number_input("H — Метеомониторинг, %",                   0.0, 100.0, float(st.session_state['H']     or 80.0), format="%.2f", key="vH")
        W_vid = st.number_input("W — Видеонаблюдение, %",                   0.0, 100.0, float(st.session_state['W_vid'] or 78.0), format="%.2f", key="vW")
        P     = st.number_input("P — Общественный транспорт, %",          -500.0, 100.0, float(st.session_state['P']     or 72.0), format="%.2f", key="vP")

    st.divider()
    ITS = 0.2*S + 0.2*Z + 0.1*M + 0.1*H + 0.2*W_vid + 0.2*P

    def _gtxt(v):
        if v >= 80:   return "🟢 Хорошо"
        elif v >= 50: return "🟡 Удовлетворительно"
        else:         return "🔴 Неудовлетворительно"

    df = pd.DataFrame({
        "Подсистема": [
            "S — Светофорное управление",
            "Z — Безопасность дорожного движения",
            "M — Мониторинг транспортного потока",
            "H — Метеомониторинг",
            "W — Видеонаблюдение и инциденты",
            "P — Общественный транспорт",
        ],
        "Значение, %":       [S, Z, M, H, W_vid, P],
        "Вес":               [0.20, 0.20, 0.10, 0.10, 0.20, 0.20],
        "Макс. вклад, %":    [20, 20, 10, 10, 20, 20],
        "Вклад в ИТСэф, %": [0.2*S, 0.2*Z, 0.1*M, 0.1*H, 0.2*W_vid, 0.2*P],
    })

    def _color_row(row):
        v = row["Значение, %"]
        bg = "#d4edda" if v >= 80 else ("#fff3cd" if v >= 50 else "#f8d7da")
        return ["background-color: " + bg if col == "Значение, %" else "" for col in row.index]

    styled_df = (
        df.style
        .apply(_color_row, axis=1)
        .format({"Значение, %": "{:.2f}", "Вес": "{:.2f}",
                 "Макс. вклад, %": "{:.0f}", "Вклад в ИТСэф, %": "{:.2f}"})
    )
    st.dataframe(styled_df, use_container_width=True, hide_index=True)

    st.markdown(
        f"**ИТСэф** = 0,2×{S:.2f} + 0,2×{Z:.2f} + 0,1×{M:.2f} "
        f"+ 0,1×{H:.2f} + 0,2×{W_vid:.2f} + 0,2×{P:.2f} = **{ITS:.2f} %**"
    )
    st.metric("🏆 ИТСэф", f"{ITS:.2f} %",
              help="Сводная формула (58) по Евстигнееву И.А. ≥ 80 % — ИТС функционирует эффективно; 50–80 % — необходима оптимизация отдельных подсистем; < 50 % — требуется комплексная модернизация.")
    if ITS >= 80:
        st.success("🟢 **Хорошо** — ИТС функционирует эффективно")
    elif ITS >= 50:
        st.warning("🟡 **Удовлетворительно** — необходима оптимизация отдельных подсистем")
    else:
        st.error("🔴 **Неудовлетворительно** — требуется комплексная модернизация ИТС")
    st.divider()

    col_radar, col_res = st.columns(2)

    with col_radar:
        st.markdown("<small style='color:rgba(49,51,63,0.6);line-height:1.2'>Диаграмма-паук показывает баллы каждой подсистемы (0–100 %).<br>Чем ближе фигура к внешней границе — тем лучше.</small>", unsafe_allow_html=True)
        labels = ['S — Светофоры', 'Z — БДД', 'M — Мониторинг',
                  'H — Метео', 'W — Видео', 'P — НГПТ']
        vals   = [max(0.0, S), max(0.0, Z), max(0.0, M),
                  max(0.0, H), max(0.0, W_vid), max(0.0, P)]
        lc = labels + [labels[0]]
        vc = vals   + [vals[0]]

        fig = go.Figure()
        fig.add_trace(go.Scatterpolar(
            r=[50]*7, theta=lc, fill='toself',
            fillcolor='rgba(255,165,0,0.08)',
            line=dict(color='orange', width=1.2, dash='dot'),
            name='50 % — удовлетворительно'
        ))
        fig.add_trace(go.Scatterpolar(
            r=[80]*7, theta=lc, fill='toself',
            fillcolor='rgba(0,180,0,0.07)',
            line=dict(color='green', width=1.2, dash='dot'),
            name='80 % — хорошо'
        ))
        fig.add_trace(go.Scatterpolar(
            r=vc, theta=lc, fill='toself',
            fillcolor='rgba(30,100,220,0.22)',
            line=dict(color='rgb(30,100,220)', width=2.5),
            mode='lines+markers',
            marker=dict(size=7),
            name='Подсистемы ИТС'
        ))
        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    range=[0, 100],
                    tickvals=[0, 25, 50, 75, 100],
                    tickfont=dict(size=10),
                    gridcolor='#ddd'
                ),
                angularaxis=dict(tickfont=dict(size=11))
            ),
            legend=dict(orientation='h', yanchor='top', y=-0.12, x=0.5,
                        xanchor='center', font=dict(size=11)),
            margin=dict(t=20, b=80, l=60, r=60),
            height=420,
        )
        st.plotly_chart(fig, use_container_width=True)

    with col_res:
        st.caption("Гистограмма показывает, сколько процентов каждая подсистема реально добавляет в итоговый ИТСэф с учётом своего веса.")
        sub_names   = ['S<br>Светофоры', 'Z<br>БДД', 'M<br>Мониторинг',
                       'H<br>Метео', 'W<br>Видео', 'P<br>НГПТ']
        contribs    = [0.2*S, 0.2*Z, 0.1*M, 0.1*H, 0.2*W_vid, 0.2*P]
        max_contrib = [20, 20, 10, 10, 20, 20]

        bar_colors = []
        for c, m in zip(contribs, max_contrib):
            ratio = c / m if m > 0 else 0
            if ratio >= 0.8:   bar_colors.append('#27ae60')
            elif ratio >= 0.5: bar_colors.append('#f39c12')
            else:              bar_colors.append('#e74c3c')

        fig_bar = go.Figure()
        fig_bar.add_trace(go.Bar(
            x=sub_names, y=max_contrib,
            marker_color='rgba(180,180,180,0.25)',
            marker_line=dict(color='#bbb', width=1),
            name='Макс. возможный вклад',
        ))
        fig_bar.add_trace(go.Bar(
            x=sub_names, y=contribs,
            marker_color=bar_colors,
            marker_line=dict(color='white', width=0.5),
            name='Фактический вклад',
            text=[f'{c:.1f} %' for c in contribs],
            textposition='outside',
            textfont=dict(size=11, color='#333'),
        ))
        for clr, lbl in [
            ('#27ae60', 'показатель ≥ 80 % — хорошо'),
            ('#f39c12', '50–80 % — удовлетворительно'),
            ('#e74c3c', '< 50 % — неудовлетворительно'),
        ]:
            fig_bar.add_trace(go.Bar(
                x=[None], y=[None],
                marker_color=clr,
                name=lbl,
                showlegend=True,
            ))
        y_min = min(min(contribs) - 2, -1)
        fig_bar.update_layout(
            barmode='overlay',
            yaxis=dict(range=[y_min, 23], title='Вклад в ИТСэф, %', gridcolor='#eee'),
            xaxis=dict(tickfont=dict(size=11)),
            legend=dict(orientation='h', y=-0.32, x=0.5, xanchor='center', font=dict(size=11)),
            plot_bgcolor='white',
            margin=dict(t=10, b=110, l=50, r=20),
            height=420,
        )
        st.plotly_chart(fig_bar, use_container_width=True)

    st.divider()

    # ── Рекомендации ─────────────────────────────────────────────
    _recs_data = [
        (S,     "S — Светофорное управление",
                "расширить охват адаптивного и диспетчерского управления; "
                "увеличить долю перекрёстков с координированным регулированием"),
        (Z,     "Z — Безопасность дорожного движения",
                "установить дополнительные средства фиксации нарушений ПДД; "
                "внедрить системы предупреждения об опасности на проблемных участках"),
        (M,     "M — Мониторинг транспортного потока",
                "увеличить число исправных датчиков и точек охвата; "
                "обеспечить регулярную передачу данных и качество аналитической обработки"),
        (H,     "H — Метеомониторинг",
                "расширить сеть АДМС; повысить точность прогнозирования "
                "(расхождение ≤ 30 % по ГОСТ Р 71094-2024) и регулярность информирования служб"),
        (W_vid, "W — Видеонаблюдение и инциденты",
                "увеличить долю работоспособных камер и охват видеоаналитикой "
                "(детектирование инцидентов, классификация ТС)"),
        (P,     "P — Общественный транспорт",
                "оптимизировать расписание; повысить пунктуальность рейсов "
                "и среднюю скорость НГПТ"),
    ]
    with st.expander("Рекомендации по развитию ИТС", expanded=True):
        need_recs = [(v, name, desc) for v, name, desc in _recs_data if v < 80]
        if not need_recs:
            st.success("Все подсистемы функционируют эффективно. "
                       "Рекомендуется поддерживать достигнутый уровень "
                       "и проводить периодический мониторинг показателей.")
        else:
            for v, name, desc in need_recs:
                icon = "🔴" if v < 50 else "🟡"
                st.markdown(f"{icon} **{name}:** {desc}.")

    st.divider()
    st.subheader("📄 Экспорт отчёта")
    stamp = datetime.now().strftime('%d%m%Y')
    col_w, col_p = st.columns(2)
    with col_w:
        word_buf = generate_word_report(S, Z, M, H, W_vid, P, ITS)
        st.download_button(
            label="⬇️ Скачать Word (.docx)",
            data=word_buf,
            file_name=f"ИТС_отчёт_{stamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with col_p:
        pdf_buf = generate_pdf_report(S, Z, M, H, W_vid, P, ITS)
        st.download_button(
            label="⬇️ Скачать PDF (.pdf)",
            data=pdf_buf,
            file_name=f"ИТС_отчёт_{stamp}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
